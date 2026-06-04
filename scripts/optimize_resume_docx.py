from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "agent_resume_source.docx"
OUT = ROOT / "agent_resume_optimized_v2.docx"

BLUE = RGBColor(0x48, 0x74, 0xCB)
BODY = RGBColor(0x40, 0x40, 0x40)


UPDATES = {
    7: [
        ("省钱智探 Agent", True, BLUE),
        ("：消费风控与记账智能体             个人项目                       2026.02 - 2026.05", True, BODY),
    ],
    10: [
        ("概述：", True, BODY),
        (
            "面向个人消费决策中“买前查价难、买后记账散、历史消费难复盘”的问题，"
            "基于 Python / Streamlit 独立开发消费风控 Agent，支持商品问价、实时比价、"
            "RAG 历史记忆、预算扣减、钉钉通知、直接记账、余额加回、撤销流水与历史持久化，"
            "形成“问价-审计-扣款-通知-复盘”的端到端闭环。",
            False,
            BODY,
        ),
    ],
    11: [
        ("Agent 流程控制", True, BLUE),
        (
            "：自研 FSM + Harness 执行层，将消费审计拆分为意图识别、实体清洗、RAG 检索、"
            "价格搜索、审计生成和账本确认等状态节点，显式控制工具调用、异常退出与结果落库，"
            "降低黑盒 Agent 执行流不可控风险。",
            False,
            BODY,
        ),
    ],
    12: [
        ("输入治理与账本闭环", True, BLUE),
        (
            "：设计规则优先的输入分流链路，支持帮助问答、商品问价、直接记账、余额加回、"
            "撤销上一笔等多类意图；基于 38 条自建消费场景回归样本验证，意图分流、"
            "金额提取和账本操作均通过测试。",
            False,
            BODY,
        ),
    ],
    13: [
        ("RAG 混合检索", True, BLUE),
        (
            "：基于 ChromaDB 记录历史购买、拦截和偏好信息，结合 Jaccard 相似度与重排序策略，"
            "将用户历史记忆、实时搜索结果和价格爬虫信息融合进审计上下文，提升口语化商品查询下的"
            "个性化判断能力。",
            False,
            BODY,
        ),
    ],
    14: [
        ("异步调度与持久化", True, BLUE),
        (
            "：使用 ThreadPoolExecutor 并行执行 RAG 检索、Web Search、价格爬虫和钉钉通知，"
            "避免外部 I/O 阻塞前端；在模拟 10 个外部通知任务测试中，主线程入队耗时约 1.47ms；"
            "使用 JSON / SQLite 持久化用户画像、历史会话和消费流水，并通过 threading.Lock 保护账本写入。",
            False,
            BODY,
        ),
    ],
    15: [
        ("稳定性与可观测", True, BLUE),
        (
            "：针对模型输出格式异常、工具调用超时、云端 Secrets 缺失、前端状态锁死等问题，"
            "设计格式解析兜底、默认价格修正、状态自动释放和异步通知失败隔离机制；构建 Agent Trace "
            "日志记录各阶段耗时与异常状态，为性能优化提供数据支撑。",
            False,
            BODY,
        ),
    ],
}

TECH_STACK_TEXT = (
    "技术栈：Python、Streamlit、OpenAI-compatible API、FSM / Harness、RAG、ChromaDB、"
    "Jaccard / Rerank、ThreadPoolExecutor、SQLite / JSON 持久化、Web Search、价格爬虫、"
    "钉钉机器人、Prompt Engineering"
)


def replace_paragraph(paragraph, parts):
    for run in paragraph.runs:
        run.text = ""
    for text, bold, color in parts:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.color.rgb = color
        run.font.name = "微软雅黑"


def main():
    doc = Document(SRC)
    for index, parts in UPDATES.items():
        replace_paragraph(doc.paragraphs[index], parts)

    tech_para = doc.paragraphs[10].insert_paragraph_before()
    replace_paragraph(tech_para, [("技术栈：", True, BLUE), (TECH_STACK_TEXT.replace("技术栈：", ""), False, BODY)])
    tech_para.style = doc.paragraphs[10].style

    for index in (8, 9):
        for run in doc.paragraphs[index].runs:
            if run.text.startswith("项目"):
                run.bold = True
                run.font.color.rgb = BLUE
            elif run.text:
                run.font.color.rgb = BODY
            run.font.size = Pt(10)
            run.font.name = "微软雅黑"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
